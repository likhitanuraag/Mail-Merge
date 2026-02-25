import sys
import re
import pandas as pd
import mammoth
import win32com.client as win32
import pythoncom
import json
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QLabel, QFileDialog, 
                             QLineEdit, QComboBox, QDialog, QTableWidget, 
                             QTableWidgetItem, QCheckBox, QSpinBox, QTextEdit, 
                             QProgressBar, QMessageBox, QGroupBox, QMenuBar, QAction)
from PyQt5.QtCore import QThread, pyqtSignal, Qt

# ==========================================
# Worker Thread for Sending Emails
# ==========================================
class MailSenderThread(QThread):
    progress_update = pyqtSignal(int, str)
    finished = pyqtSignal(bool, str)

    def __init__(self, data_df, template_html, subject_template, mapping, cc_col, bcc_col, 
                 email_col, send_as_draft, start_row, end_row):
        super().__init__()
        self.data_df = data_df
        self.template_html = template_html
        self.subject_template = subject_template
        self.mapping = mapping
        self.cc_col = cc_col
        self.bcc_col = bcc_col
        self.email_col = email_col
        self.send_as_draft = send_as_draft
        self.start_row = start_row
        self.end_row = end_row

    def run(self):
        try:
            pythoncom.CoInitialize()
            outlook = win32.Dispatch('outlook.application')
            
            total_records = self.end_row - self.start_row
            success_count = 0
            failed_records = []
            
            for index in range(self.start_row, self.end_row):
                row = self.data_df.iloc[index]
                recipient_email = str(row[self.email_col]) if self.email_col and pd.notna(row[self.email_col]) else "Unknown/Empty"
                
                try:
                    subject = self.subject_template
                    body_html = self.template_html
                    
                    for placeholder, col_name in self.mapping.items():
                        val = str(row[col_name]) if pd.notna(row[col_name]) else ""
                        subject = subject.replace(f"{{{{{placeholder}}}}}", val)
                        body_html = body_html.replace(f"{{{{{placeholder}}}}}", val)

                    mail = outlook.CreateItem(0)
                    mail.To = recipient_email
                    
                    if self.cc_col and pd.notna(row[self.cc_col]):
                        mail.CC = str(row[self.cc_col])
                    if self.bcc_col and pd.notna(row[self.bcc_col]):
                        mail.BCC = str(row[self.bcc_col])
                        
                    mail.Subject = subject
                    mail.HTMLBody = body_html 
                    
                    if self.send_as_draft:
                        mail.Save()
                    else:
                        mail.Send()
                        
                    success_count += 1
                    status_msg = f"Processed {index - self.start_row + 1}/{total_records}: {recipient_email}"
                    
                except Exception as e:
                    error_msg = str(e)
                    failed_records.append((index + 1, recipient_email, error_msg))
                    status_msg = f"FAILED {index - self.start_row + 1}/{total_records}: {recipient_email}"
                
                progress_pct = int(((index - self.start_row + 1) / total_records) * 100)
                self.progress_update.emit(progress_pct, status_msg)
                
            if not failed_records:
                final_msg = f"Successfully processed all {success_count} emails!"
                self.finished.emit(True, final_msg)
            else:
                final_msg = f"Processed {success_count} successfully, but {len(failed_records)} failed.\n\nFailures:\n"
                for r_idx, email, err in failed_records:
                    final_msg += f"- Row {r_idx} ({email}): {err}\n"
                self.finished.emit(False, final_msg)
            
        except Exception as e:
            self.finished.emit(False, f"FATAL ERROR:\n{str(e)}")
            
        finally:
            pythoncom.CoUninitialize()

# ==========================================
# Help / SOP Dialog
# ==========================================
class HelpDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("SOP - Mail Merge Utility")
        self.resize(650, 600)
        
        layout = QVBoxLayout(self)
        
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        
        sop_html = """
        <h2 style='color: #2e6c80;'>SOP - Mail Merge Utility</h2>
        <p>Follow these steps to execute a successful merge:</p>

        <h3>Step 1: Prepare Your Documents</h3>
        <ul>
            <li><strong>Word Document (.docx):</strong> Type <code>{{PlaceholderName}}</code> wherever you want dynamic data to appear (e.g., <code>Dear {{First Name}},</code>). Ensure you do not change font formatting mid-placeholder.</li>
            <li><strong>Excel Data (.xlsx):</strong> Ensure the first row contains clear column headers.</li>
        </ul>

        <h3>Step 2: Load Files</h3>
        <ul>
            <li>Click <strong>Browse Word Document</strong> and select your template.</li>
            <li>Click <strong>Browse Excel Data</strong> and select your mailing list.</li>
        </ul>

        <h3>Step 3: Map Columns & Settings</h3>
        <ul>
            <li>Click <strong>Map Columns</strong>. Match each Word placeholder to the corresponding Excel column.</li>
            <li>Select the Email column for the <strong>To</strong> field. Optionally, select columns for <strong>CC</strong> and <strong>BCC</strong>.</li>
            <li>Enter a <strong>Subject Line</strong>. You can use exact Excel header names inside brackets (e.g., <code>Invoice for {{Company Name}}</code>) to make subjects dynamic.</li>
            <li><em>Tip:</em> Use <strong>Save Config</strong> to save these mappings to a .json file so you can load them instantly next time!</li>
        </ul>

        <h3>Step 4: Preview</h3>
        <ul>
            <li>Use the <strong>&lt; Prev</strong> and <strong>Next &gt;</strong> buttons to cycle through records.</li>
            <li>Review the HTML rendering, To/CC/BCC routing, and Subject line.</li>
            <li>If you edit your Word doc or Excel file externally, click <strong>Refresh Preview</strong> to reload the data without restarting the app.</li>
        </ul>

        <h3>Step 5: Process Emails</h3>
        <ul>
            <li>Select your row range (default is all rows).</li>
            <li>Leave <strong>Save as Drafts</strong> checked to push the emails to your Outlook Drafts folder for final review. Uncheck it only when you are ready to send live immediately.</li>
            <li>Click <strong>Process Emails</strong> and wait for the success dialogue.</li>
        </ul>
        
        <p>https://github.com/likhitanuraag</p>
        """
        text_edit.setHtml(sop_html)
        
        layout.addWidget(text_edit)
        
        close_btn = QPushButton("Close Help")
        close_btn.setStyleSheet("padding: 8px; font-weight: bold;")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

# ==========================================
# Mapping Modal Dialog
# ==========================================
class MappingDialog(QDialog):
    def __init__(self, placeholders, excel_columns, current_mapping, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Map Word Placeholders to Excel Columns")
        self.resize(450, 400)
        self.mapping = current_mapping.copy()
        
        layout = QVBoxLayout(self)
        
        self.table = QTableWidget(len(placeholders), 2)
        self.table.setHorizontalHeaderLabels(["Word Placeholder", "Excel Column"])
        self.table.horizontalHeader().setStretchLastSection(True)
        
        self.combos = []
        for i, ph in enumerate(placeholders):
            self.table.setItem(i, 0, QTableWidgetItem(f"{{{{{ph}}}}}"))
            combo = QComboBox()
            combo.addItem("-- Ignore --")
            combo.addItems(excel_columns)
            
            # Auto-match or load existing mapping
            if ph in self.mapping and self.mapping[ph] in excel_columns:
                combo.setCurrentText(self.mapping[ph])
            elif ph in excel_columns:
                combo.setCurrentText(ph)
                
            self.table.setCellWidget(i, 1, combo)
            self.combos.append((ph, combo))
            
        layout.addWidget(self.table)
        
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("Save Mapping")
        save_btn.clicked.connect(self.save_mapping)
        btn_layout.addWidget(save_btn)
        layout.addLayout(btn_layout)
        
    def save_mapping(self):
        self.mapping.clear()
        for ph, combo in self.combos:
            if combo.currentText() != "-- Ignore --":
                self.mapping[ph] = combo.currentText()
        self.accept()

# ==========================================
# Main Application Window
# ==========================================
class MailMergeApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Mail Merge Utility")
        self.resize(900, 800)
        
        self.word_path = ""
        self.excel_path = ""
        self.df = pd.DataFrame()
        self.template_html = ""
        self.placeholders = []
        self.mapping = {}
        
        self.init_ui()
        self.create_menu()
        
    def create_menu(self):
        menubar = self.menuBar()
        help_menu = menubar.addMenu("Help")
        
        sop_action = QAction("How to Use (SOP)", self)
        sop_action.triggered.connect(self.show_help)
        help_menu.addAction(sop_action)

        sample_action = QAction("Generate Sample Files", self)
        sample_action.triggered.connect(self.generate_samples)
        help_menu.addAction(sample_action)
        
    def show_help(self):
        dialog = HelpDialog(self)
        dialog.exec_()

    def generate_samples(self):
        try:
            import docx
            import openpyxl 
        except ImportError:
            QMessageBox.warning(self, "Missing Dependencies", 
                                "To generate sample files, please install python-docx and openpyxl.\n\nRun this in your terminal:\npip install python-docx openpyxl")
            return

        folder = QFileDialog.getExistingDirectory(self, "Select Folder to Save Samples")
        if not folder:
            return

        try:
            # 1. Generate Generic Sample Excel Data
            df = pd.DataFrame({
                "First Name": ["Jane", "John", "Alice"],
                "Last Name": ["Smith", "Doe", "Johnson"],
                "Email": ["jane.smith@example.com", "john.doe@example.com", "alice.j@example.com"],
                "CC Email": ["billing.cc@example.com", "", ""],
                "Account Manager": ["Sarah", "Michael", "Sarah"],
                "Invoice Number": ["INV-1001", "INV-1002", "INV-1003"],
                "Service Date": ["2026-01-15", "2026-01-18", "2026-01-20"],
                "Description": ["Web Hosting Services", "Annual Consulting", "Software Licensing"],
                "Amount Due": ["$150.00", "$450.00", "$99.00"],
                "Due Date": ["2026-02-15", "2026-02-18", "2026-02-20"]
            })
            excel_path = f"{folder}/Sample_Data.xlsx"
            df.to_excel(excel_path, index=False)

            # 2. Generate Generic Sample Word Document
            doc = docx.Document()
            doc.add_paragraph('Dear {{First Name}} {{Last Name}},')
            doc.add_paragraph('Thank you for your continued business with Acme Corporation. This is an automated notification regarding your recent account activity.')
            doc.add_paragraph('Below are the details of your latest invoice:')
            
            # Add a generic table
            table = doc.add_table(rows=2, cols=5)
            table.style = 'Table Grid'
            
            headers = ['Invoice Number', 'Service Date', 'Description', 'Amount Due', 'Due Date']
            placeholders = ['{{Invoice Number}}', '{{Service Date}}', '{{Description}}', '{{Amount Due}}', '{{Due Date}}']
            
            for i in range(5):
                table.rows[0].cells[i].text = headers[i]
                table.rows[1].cells[i].text = placeholders[i]
            
            doc.add_paragraph('\nPlease ensure payment is completed by {{Due Date}}. If you have any questions, your account manager, {{Account Manager}}, will be happy to assist you.')
            
            doc.add_paragraph('\nBest Regards,')
            doc.add_paragraph('Billing Department')
            doc.add_paragraph('Acme Corporation')
            doc.add_paragraph('Email: billing@acmecorp.example.com')
            doc.add_paragraph('Phone: +1 800-555-0199')
            
            word_path = f"{folder}/Sample_Template.docx"
            doc.save(word_path)

            QMessageBox.information(self, "Success", f"Sample files successfully created in:\n{folder}\n\n• Sample_Data.xlsx\n• Sample_Template.docx")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate samples:\n{str(e)}")


    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        
        # --- File Selection ---
        file_group = QGroupBox("1. Load Files")
        file_layout = QVBoxLayout()
        
        self.lbl_word = QLabel("Word Doc: Not selected")
        btn_word = QPushButton("Browse Word Document")
        btn_word.clicked.connect(self.prompt_load_word)
        
        self.lbl_excel = QLabel("Excel File: Not selected")
        btn_excel = QPushButton("Browse Excel Data")
        btn_excel.clicked.connect(self.prompt_load_excel)
        
        file_layout.addWidget(self.lbl_word)
        file_layout.addWidget(btn_word)
        file_layout.addWidget(self.lbl_excel)
        file_layout.addWidget(btn_excel)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        # --- Mapping & Settings ---
        settings_group = QGroupBox("2. Settings & Mapping")
        settings_layout = QVBoxLayout()
        
        map_btn_layout = QHBoxLayout()
        self.btn_map = QPushButton("Map Columns")
        self.btn_map.clicked.connect(self.open_mapping)
        self.btn_map.setEnabled(False)
        
        self.btn_save_config = QPushButton("Save Config")
        self.btn_save_config.clicked.connect(self.save_config)
        
        self.btn_load_config = QPushButton("Load Config")
        self.btn_load_config.clicked.connect(self.load_config)
        
        map_btn_layout.addWidget(self.btn_map)
        map_btn_layout.addWidget(self.btn_save_config)
        map_btn_layout.addWidget(self.btn_load_config)
        settings_layout.addLayout(map_btn_layout)
        
        email_layout = QHBoxLayout()
        email_layout.addWidget(QLabel("To (Email Column):"))
        self.combo_to = QComboBox()
        self.combo_to.currentTextChanged.connect(lambda: self.update_preview(0))
        email_layout.addWidget(self.combo_to)
        
        email_layout.addWidget(QLabel("CC Column:"))
        self.combo_cc = QComboBox()
        self.combo_cc.currentTextChanged.connect(lambda: self.update_preview(0))
        email_layout.addWidget(self.combo_cc)
        
        email_layout.addWidget(QLabel("BCC Column:"))
        self.combo_bcc = QComboBox()
        self.combo_bcc.currentTextChanged.connect(lambda: self.update_preview(0))
        email_layout.addWidget(self.combo_bcc)
        settings_layout.addLayout(email_layout)
        
        settings_layout.addWidget(QLabel("Subject Line (Use {{column}} for placeholders):"))
        self.txt_subject = QLineEdit()
        self.txt_subject.textChanged.connect(lambda: self.update_preview(0))
        settings_layout.addWidget(self.txt_subject)
        
        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)
        
        # --- Preview ---
        preview_group = QGroupBox("3. HTML Preview")
        preview_layout = QVBoxLayout()
        
        nav_layout = QHBoxLayout()
        self.btn_prev = QPushButton("< Prev")
        self.btn_prev.clicked.connect(lambda: self.update_preview(-1))
        
        self.lbl_record = QLabel("Record: 0/0")
        self.lbl_record.setAlignment(Qt.AlignCenter)
        
        self.btn_next = QPushButton("Next >")
        self.btn_next.clicked.connect(lambda: self.update_preview(1))
        
        self.btn_refresh = QPushButton("Refresh Preview")
        self.btn_refresh.clicked.connect(self.refresh_preview)
        
        nav_layout.addWidget(self.btn_prev)
        nav_layout.addWidget(self.lbl_record)
        nav_layout.addWidget(self.btn_next)
        nav_layout.addWidget(self.btn_refresh)
        
        self.txt_preview = QTextEdit()
        self.txt_preview.setReadOnly(True)
        self.txt_preview.setStyleSheet("background-color: #ffffff; color: #000000;") 
        
        preview_layout.addLayout(nav_layout)
        preview_layout.addWidget(self.txt_preview)
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # --- Sending Options ---
        send_group = QGroupBox("4. Send Options")
        send_layout = QVBoxLayout()
        
        opts_layout = QHBoxLayout()
        self.chk_draft = QCheckBox("Save as Drafts (Do not send immediately)")
        self.chk_draft.setChecked(True)
        opts_layout.addWidget(self.chk_draft)
        
        opts_layout.addWidget(QLabel("Start Row:"))
        self.spin_start = QSpinBox()
        self.spin_start.setMinimum(1)
        opts_layout.addWidget(self.spin_start)
        
        opts_layout.addWidget(QLabel("End Row:"))
        self.spin_end = QSpinBox()
        self.spin_end.setMinimum(1)
        opts_layout.addWidget(self.spin_end)
        
        send_layout.addLayout(opts_layout)
        
        self.btn_send = QPushButton("Process Emails")
        self.btn_send.clicked.connect(self.process_emails)
        self.btn_send.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 10px;")
        send_layout.addWidget(self.btn_send)
        
        self.lbl_status = QLabel("Status: Waiting...")
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        
        send_layout.addWidget(self.lbl_status)
        send_layout.addWidget(self.progress_bar)
        
        send_group.setLayout(send_layout)
        layout.addWidget(send_group)
        
        self.current_preview_index = 0

    # --- Core Loading Functions ---
    def prompt_load_word(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Word Document", "", "Word Files (*.docx)")
        if path:
            self._process_word(path)

    def _process_word(self, path):
        try:
            with open(path, "rb") as docx_file:
                raw_text = mammoth.extract_raw_text(docx_file).value
                self.placeholders = list(set(re.findall(r'\{\{(.*?)\}\}', raw_text)))
                
                docx_file.seek(0)
                result = mammoth.convert_to_html(docx_file)
                
                # INJECT CSS FOR TABLE STYLING
                table_css = """
                <style>
                    table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
                    th, td { border: 1px solid #999999; padding: 8px; text-align: left; }
                    th { background-color: #f2f2f2; }
                </style>
                """
                self.template_html = table_css + result.value
                
            self.word_path = path
            self.lbl_word.setText(f"Word Doc: {path.split('/')[-1]}")
            self.check_ready()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not read Word document:\n{str(e)}")

    def prompt_load_excel(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Excel Data", "", "Excel Files (*.xlsx *.xls)")
        if path:
            self._process_excel(path)

    def _process_excel(self, path):
        try:
            self.df = pd.read_excel(path)
            columns = self.df.columns.tolist()
            
            for combo in [self.combo_to, self.combo_cc, self.combo_bcc]:
                combo.blockSignals(True)
                current = combo.currentText()
                combo.clear()
                combo.addItem("-- None --")
                combo.addItems(columns)
                if current in columns:
                    combo.setCurrentText(current)
                elif combo == self.combo_to:
                    for col in columns:
                        if 'email' in col.lower() or 'mail' in col.lower():
                            combo.setCurrentText(col)
                            break
                combo.blockSignals(False)
                            
            self.spin_start.setMaximum(len(self.df))
            self.spin_end.setMaximum(len(self.df))
            self.spin_end.setValue(len(self.df))
            
            self.excel_path = path
            self.lbl_excel.setText(f"Excel File: {path.split('/')[-1]}")
            self.check_ready()
            self.update_preview(0)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not read Excel file:\n{str(e)}")

    def check_ready(self):
        if self.word_path and self.excel_path:
            self.btn_map.setEnabled(True)

    def open_mapping(self):
        dialog = MappingDialog(self.placeholders, self.df.columns.tolist(), self.mapping, self)
        if dialog.exec_() == QDialog.Accepted:
            self.mapping = dialog.mapping
            self.current_preview_index = 0
            self.update_preview(0)

    # --- Save / Load Config ---
    def save_config(self):
        if not self.mapping:
            QMessageBox.warning(self, "Warning", "Please map your columns first.")
            return
            
        path, _ = QFileDialog.getSaveFileName(self, "Save Configuration", "", "JSON Files (*.json)")
        if path:
            config_data = {
                "mapping": self.mapping,
                "subject": self.txt_subject.text(),
                "to": self.combo_to.currentText(),
                "cc": self.combo_cc.currentText(),
                "bcc": self.combo_bcc.currentText()
            }
            try:
                with open(path, 'w') as f:
                    json.dump(config_data, f, indent=4)
                QMessageBox.information(self, "Success", "Configuration saved successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save config:\n{str(e)}")

    def load_config(self):
        path, _ = QFileDialog.getOpenFileName(self, "Load Configuration", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, 'r') as f:
                    config_data = json.load(f)
                    
                self.mapping = config_data.get("mapping", {})
                self.txt_subject.blockSignals(True)
                self.txt_subject.setText(config_data.get("subject", ""))
                self.txt_subject.blockSignals(False)
                
                for combo, key in [(self.combo_to, "to"), (self.combo_cc, "cc"), (self.combo_bcc, "bcc")]:
                    combo.blockSignals(True)
                    val = config_data.get(key, "-- None --")
                    idx = combo.findText(val)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)
                    combo.blockSignals(False)
                        
                self.update_preview(0)
                QMessageBox.information(self, "Success", "Configuration loaded successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load config:\n{str(e)}")

    # --- Preview & Refresh ---
    def refresh_preview(self):
        if self.word_path:
            self._process_word(self.word_path)
        if self.excel_path:
            self._process_excel(self.excel_path)
        self.update_preview(0)

    def update_preview(self, step):
        if self.df.empty or not self.mapping:
            return
            
        self.current_preview_index += step
        self.current_preview_index = max(0, min(self.current_preview_index, len(self.df) - 1))
        
        row = self.df.iloc[self.current_preview_index]
        
        preview_body_html = self.template_html
        preview_subject = self.txt_subject.text()
        
        for placeholder, col_name in self.mapping.items():
            val = str(row[col_name]) if pd.notna(row[col_name]) else ""
            preview_body_html = preview_body_html.replace(f"{{{{{placeholder}}}}}", val)
            preview_subject = preview_subject.replace(f"{{{{{placeholder}}}}}", val)
            
        subject_phs = set(re.findall(r'\{\{(.*?)\}\}', preview_subject))
        for ph in subject_phs:
            if ph in self.df.columns:
                val = str(row[ph]) if pd.notna(row[ph]) else ""
                preview_subject = preview_subject.replace(f"{{{{{ph}}}}}", val)
                
        to_col = self.combo_to.currentText()
        to_email = str(row[to_col]) if to_col != "-- None --" and pd.notna(row[to_col]) else "No Email Column Mapped!"
        
        cc_col = self.combo_cc.currentText()
        cc_email = str(row[cc_col]) if cc_col != "-- None --" and pd.notna(row[cc_col]) else ""
        cc_html = f"<strong>CC:</strong> {cc_email}<br>" if cc_email else ""
        
        bcc_col = self.combo_bcc.currentText()
        bcc_email = str(row[bcc_col]) if bcc_col != "-- None --" and pd.notna(row[bcc_col]) else ""
        bcc_html = f"<strong>BCC:</strong> {bcc_email}<br>" if bcc_email else ""
            
        header_html = f"""
        <div style='font-family: Arial, sans-serif; margin-bottom: 20px; border-bottom: 2px solid #ccc; padding-bottom: 10px;'>
            <strong>TO:</strong> {to_email}<br>
            {cc_html}
            {bcc_html}
            <strong>SUBJECT:</strong> {preview_subject}
        </div>
        """
        
        self.txt_preview.setHtml(header_html + preview_body_html)
        self.lbl_record.setText(f"Record: {self.current_preview_index + 1}/{len(self.df)}")

    # --- Sending ---
    def process_emails(self):
        if self.combo_to.currentText() == "-- None --":
            QMessageBox.warning(self, "Error", "Please select an Email column for the 'To' field.")
            return
            
        self.btn_send.setEnabled(False)
        self.progress_bar.setValue(0)
        
        self.thread = MailSenderThread(
            data_df=self.df,
            template_html=self.template_html,
            subject_template=self.txt_subject.text(),
            mapping=self.mapping,
            cc_col=self.combo_cc.currentText() if self.combo_cc.currentText() != "-- None --" else None,
            bcc_col=self.combo_bcc.currentText() if self.combo_bcc.currentText() != "-- None --" else None,
            email_col=self.combo_to.currentText(),
            send_as_draft=self.chk_draft.isChecked(),
            start_row=self.spin_start.value() - 1,
            end_row=self.spin_end.value()
        )
        
        self.thread.progress_update.connect(self.update_progress)
        self.thread.finished.connect(self.thread_finished)
        self.thread.start()

    def update_progress(self, val, msg):
        self.progress_bar.setValue(val)
        self.lbl_status.setText(f"Status: {msg}")

    def thread_finished(self, completely_successful, msg):
        self.btn_send.setEnabled(True)
        
        box = QMessageBox(self)
        if completely_successful:
            box.setIcon(QMessageBox.Information)
            box.setWindowTitle("Merge Complete")
            box.setText("Process Completed Successfully!")
            box.setInformativeText(msg)
        else:
            if "FATAL ERROR:" in msg:
                box.setIcon(QMessageBox.Critical)
                box.setWindowTitle("Fatal Error")
                box.setText("The process crashed before completing.")
                box.setInformativeText(msg)
            else:
                box.setIcon(QMessageBox.Warning)
                box.setWindowTitle("Completed with Errors")
                box.setText("The mail merge finished, but some emails failed.")
                
                parts = msg.split("\n\nFailures:\n")
                box.setInformativeText(parts[0])
                if len(parts) > 1:
                    box.setDetailedText("Failures:\n" + parts[1])
                    
        box.exec_()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MailMergeApp()
    window.show()
    sys.exit(app.exec_())